from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from ..core.models import XMindDocument, Sheet, TopicNode
from ..utils.file_manager import FileManager
from ..utils.image_handler import ImageHandler
from .base import BaseExporter


class DocxExporter(BaseExporter):
    """Word (.docx) 格式导出器"""
    
    def export(self, output_path: Path, sheet_index: int = 0):
        """导出单个sheet为Word"""
        sheet = self._get_sheet(sheet_index)
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        
        # 创建图片处理器 - 图片保存在 output/文件名/images/ 目录
        images_dir = output_path.parent.parent / 'images'
        image_handler = ImageHandler(images_dir)
        
        # 添加内容
        self._add_node_to_doc(doc, sheet.root_topic, image_handler, level=0)
        
        # 保存文件
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
    
    def export_all_sheets(self, output_dir: Path):
        """导出所有sheet为Word"""
        output_dir.mkdir(parents=True, exist_ok=True)
        
        for i, sheet in enumerate(self.document.sheets):
            filename = f"{sheet.title or f'Sheet_{i+1}'}.docx"
            output_path = output_dir / filename
            self.export(output_path, sheet_index=i)
    
    def _add_node_to_doc(self, doc: Document, node: TopicNode, 
                         image_handler: ImageHandler, level: int):
        """递归添加节点到 Word 文档"""
        if not node or not node.title:
            return
        
        title = node.title
        
        # 添加标题
        if level == 0:
            heading = doc.add_heading(title, level=0)
        else:
            heading_level = min(level, 4)
            heading = doc.add_heading(title, level=heading_level)
        
        # 添加备注
        if node.notes:
            p = doc.add_paragraph()
            p.add_run(f"备注: {node.notes}").italic = True
        
        # 添加标签
        if node.labels:
            labels_text = ', '.join(node.labels)
            p = doc.add_paragraph()
            p.add_run(f"标签: {labels_text}").bold = True
        
        # 添加图片
        for img_path_str in node.images:
            img_path = Path(img_path_str)
            if img_path.exists():
                try:
                    # 添加图片到 Word
                    doc.add_picture(str(img_path), width=Inches(5.0))
                except Exception as e:
                    print(f"警告: 无法添加图片 {img_path}: {e}")
        
        # 添加空行
        if level == 0 or node.children:
            doc.add_paragraph()
        
        # 递归处理子节点
        for child in node.children:
            self._add_node_to_doc(doc, child, image_handler, level + 1)
