from pathlib import Path
from docx import Document
from ..core.models import XMindDocument, Sheet, TopicNode
from .base import BaseImporter


class DocxImporter(BaseImporter):
    """Word 文档导入器（Word -> XMind）"""
    
    def parse(self) -> XMindDocument:
        """解析 Word 文档"""
        doc = Document(self.file_path)
        
        # 解析 Word 内容
        root_node = self._parse_docx_content(doc)
        
        # 创建 Sheet
        sheet_title = self.file_path.stem
        sheet = self._create_sheet(sheet_title, root_node)
        self.document.add_sheet(sheet)
        
        return self.document
    
    def _parse_docx_content(self, doc: Document) -> TopicNode:
        """解析 Word 内容，构建树结构"""
        # 收集标题和正文
        headings = []  # (level, title, body_paragraphs)
        current_body = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                # 空行也保留作为正文的一部分
                if current_body:
                    current_body.append('')
                continue
            
            # 检查是否是标题
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                    
                    # 保存前一个标题的正文
                    if headings:
                        headings[-1] = (headings[-1][0], headings[-1][1], current_body)
                    
                    headings.append((level, text, []))
                    current_body = []
                except ValueError:
                    # 不是标准标题格式，当作正文
                    current_body.append(text)
            else:
                # 非标题段落，保存为正文
                current_body.append(text)
        
        # 保存最后一个标题的正文
        if headings and current_body:
            headings[-1] = (headings[-1][0], headings[-1][1], current_body)
        
        # 创建根节点
        if headings:
            root_level, root_title, root_body = headings[0]
            root = TopicNode(title=root_title, level=0)
            if root_body:
                root.notes = '\n'.join(root_body).strip()
            child_headings = headings[1:]
        else:
            # 没有标题，使用第一段作为根节点
            root_title = doc.paragraphs[0].text.strip() if doc.paragraphs else self.file_path.stem
            root = TopicNode(title=root_title, level=0)
            child_headings = []
        
        # 构建树结构
        if child_headings:
            self._build_tree_with_body(root, child_headings)
        
        return root
    
    def _build_tree_with_body(self, parent: TopicNode, headings: list):
        """根据节点列表构建树结构，正文作为子节点添加"""
        if not headings:
            return
        
        node_stack = [(parent, 0)]  # (node, level)
        
        for level, title, body_paras in headings:
            # 创建标题节点
            new_node = TopicNode(title=title, level=level)
            
            # 添加正文作为子节点（如果有）
            if body_paras:
                body_text = '\n'.join(body_paras).strip()
                if body_text:
                    # 将正文内容直接作为子节点标题
                    body_node = TopicNode(title=body_text, level=level + 1)
                    new_node.add_child(body_node)
            
            # 找到正确的父节点
            while node_stack and node_stack[-1][1] >= level:
                node_stack.pop()
            
            if node_stack:
                parent_node = node_stack[-1][0]
                parent_node.add_child(new_node)
            else:
                parent.add_child(new_node)
            
            # 将新节点压入栈
            node_stack.append((new_node, level))
